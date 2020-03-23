import docx
import pdb
import glob
import os
import re
from datetime import datetime

## SET THIS BEFORE USING PROGRAM!
starting_hours_for_overtime = 8


def main():
  print("Starting program")

  # collect files
  os.chdir('.')
  for attendance_sheet in glob.glob('*.docx'):
    print(f"*** Starting to parse {attendance_sheet} ***")

    # opening file
    file_name = attendance_sheet
    doc = docx.Document(file_name)
    
    tables = doc.tables
    date_range = doc.paragraphs[2].text.replace(u'\xa0', u' ')
    file_name = re.sub("[^0-9\-]", "", date_range)
    employees = []
    for table in tables:
      employees.append(parse_employee_table(table))

    for employee in employees:
      calculate_hours(employee, starting_hours_for_overtime)
    write_to_file(f"{file_name}_docx", employees)
  
  print("Ending program")

def parse_employee_table(table):
  employee = {}
  employee_name = table.cell(0,1).text.replace(u'\xa0', u' ').split()[1].split(':')[1]
  employee['name'] = employee_name
  employee['attendance'] = {}

  column_length = len(table.columns)
  row_length = len(table.rows)
  data = []
  for y in range(0, int((row_length - 1) / 2)):
    for x in range(1, column_length):
      dates_row_index = 2*y+1
      times_row_index = (2*(y+1))
      date = table.cell(dates_row_index, x).text.replace(u'\xa0', u' ').strip()
      time_contents = table.cell(times_row_index, x).text.replace(u'\xa0', u' ').strip().split('\n')
      time_contents = [times for times in time_contents if (isinstance(times, str) and times != '')]
      if(date == ''):
        continue
      else:
        employee['attendance'][date] = time_contents
  return employee


# employee and hours until counts of overtime, defaults at 8
# ASSUMES EMPLOYEES CLOCK IN AND OUT ON THE SAME DATE
def calculate_hours(employee, overtime_hours_min=8):
  attendance = employee['attendance']
  attendance = attendance.items()
  review = {}
  num_hours_worked_total = 0
  num_overtime_hours_worked = 0
  num_regular_hours_worked = 0
  for day in attendance:
    date = day[0]
    log = day[1]
    # error handling if didn't work or forgot to sign in/out
    if(len(log) == 0):
      continue
    if(len(log) == 1):
      print(f"{employee['name']} forgot to sign-in or out on {date}")
      review[date] = log
      continue
    # start calculating hours
    start_time = datetime.strptime(log[0], '%H:%M')
    end_time = datetime.strptime(log[1], '%H:%M')
    hours_worked = (end_time - start_time).total_seconds() / 3600
    overtime_hours_worked = 0
    if(hours_worked - overtime_hours_min > 0):
      overtime_hours_worked = hours_worked - overtime_hours_min
      regular_hours_worked = overtime_hours_min
    else:
      regular_hours_worked = hours_worked
    num_hours_worked_total += hours_worked
    num_overtime_hours_worked += overtime_hours_worked
    num_regular_hours_worked += regular_hours_worked
  employee["review"] = review
  employee["num_hours_worked_total"] = num_hours_worked_total
  employee["num_overtime_hours_worked"] = num_overtime_hours_worked
  employee["num_regular_hours_worked"] = num_regular_hours_worked


def write_to_file(file_name, employees):
  file = open(f"./{file_name}.txt", "w", encoding="utf-8")
  for employee in employees:
    file.write(f"Employee name: {employee['name']}\n")
    file.write(
      f"Total hours worked: {employee['num_hours_worked_total']}\n")
    file.write(
      f"Overtime hours worked: {employee['num_overtime_hours_worked']}\n")
    file.write(
      f"Regular hours worked: {employee['num_regular_hours_worked']}\n")
    if(len(employee["review"].items()) > 0):
      file.write(f"---------REQUIRES REVIEW!---------\n")
      for review in employee["review"].items():
        file.write(f"Date: {review[0]}, Time: {review[1]}\n")
      file.write(f"---------REQUIRES REVIEW!---------\n")
    file.write(f"\n\n")
  print(f"COMPLETE: {file_name}.txt")

if __name__ == "__main__":
  main()
